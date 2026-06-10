# ═══════════════════════════════════════════════════════════════════════
# Copyright (c) 2025 Pratyush Chaudhari. All rights reserved.
#
# This source code is part of the Ethical Robustness Testing System (ERTS).
# Research paper: https://zenodo.org/records/20544025
#
# LICENSE: This code is provided for academic study and personal
# learning ONLY. Commercial use, corporate deployment, or any use
# intended to generate revenue is strictly prohibited without
# explicit written permission from the author.
# ═══════════════════════════════════════════════════════════════════════

"""Generate IEEE-format research paper for ERTS as a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0

# Heading styles
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)

doc.styles['Heading 1'].font.size = Pt(11)
doc.styles['Heading 2'].font.size = Pt(10)
doc.styles['Heading 3'].font.size = Pt(10)

def add_centered(text, size=10, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Cambria Math'
    r.font.size = Pt(10)
    r.italic = True
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════
add_centered('System and Method for Adversarial Evaluation of Ethical AI Decision-Making Models Using Semantic Perturbation Functions in a Bounded Ethical Consequence Space with Domain-Adaptive Deployment Certification', size=14, bold=True)
doc.add_paragraph()
add_centered('Pratyush', size=11)
add_centered('Independent Researcher', size=10, italic=True)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('Abstract')
r.bold = True
r.italic = True
r.font.size = Pt(9)
p.add_run(' -- ').bold = True

abstract_text = (
    'As artificial intelligence systems are increasingly deployed in high-stakes ethical decision-making contexts '
    'such as healthcare triage, autonomous vehicle control, and employment screening, there exists a critical gap '
    'in formal methods for evaluating their robustness against adversarial manipulation of ethical reasoning. '
    'Existing adversarial testing frameworks operate on raw data representations (pixels, tokens, features) and '
    'cannot evaluate the semantic stability of ethical judgments under structured pressure. '
    'This paper introduces the Ethical Robustness Testing System (ERTS), a closed-pipeline framework that: '
    '(1) encodes ethical dilemmas into a 22-dimensional Ethical Consequence Space (ECS) where each dimension '
    'represents a named ethical variable with semantic meaning; '
    '(2) applies 17 formal semantic perturbation functions across 7 adversarial categories, subject to 6 validity '
    'constraint classes including a novel semantic coherence constraint; '
    '(3) measures decision deviation via a 4-component Ethical Instability Index (EII); and '
    '(4) produces domain-adaptive deployment certification verdicts across 8 application domains. '
    'We evaluate 4 structurally distinct moral AI architectures (rule-based, learning-based, RLHF-aligned, and '
    'virtue ethics) across 20 ethical scenarios spanning 6 real-world categories, generating 400 adversarial '
    'test cases. Results demonstrate that only 25% of models achieve deployment certification even under '
    'baseline thresholds, revealing critical vulnerabilities in fairness preservation and authority resistance. '
    'The ECS representation, the EII metric, and the semantic coherence constraint system constitute novel '
    'contributions with no identified prior art in the adversarial machine learning or AI ethics literature.'
)
r2 = p.add_run(abstract_text)
r2.font.size = Pt(9)
r2.italic = True

doc.add_paragraph()

# Keywords
p = doc.add_paragraph()
r = p.add_run('Index Terms')
r.bold = True
r.italic = True
r.font.size = Pt(9)
p.add_run(' -- ').bold = True
r2 = p.add_run('Ethical AI, Adversarial Robustness, Semantic Perturbation, Ethical Consequence Space, '
               'AI Safety Certification, Machine Ethics, Ethical Instability Index, Domain-Adaptive Testing')
r2.font.size = Pt(9)
r2.italic = True

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('I. INTRODUCTION', level=1)

add_body(
    'The deployment of AI systems in ethical decision-making contexts has accelerated significantly, '
    'with applications ranging from clinical decision support [1] to autonomous vehicle control [2], '
    'algorithmic hiring [3], and military target identification [4]. These systems are expected to make '
    'decisions that are not only accurate but ethically sound, fair, and resistant to manipulation [5].'
)
add_body(
    'While adversarial robustness testing has become standard practice for evaluating the reliability '
    'of machine learning models [6], [7], existing frameworks operate exclusively on raw data '
    'representations. The Adversarial Robustness Toolbox (ART) [8] applies perturbations to pixel '
    'values and input features. NVIDIA Garak [9] performs red-teaming on text generation models. '
    'TrustLLM [10] and Stanford HELM [11] benchmark safety properties of large language models. '
    'However, none of these systems can evaluate whether an AI model\'s ethical judgment remains '
    'stable when the ethical framing of a scenario is deliberately manipulated.'
)
add_body(
    'This gap is consequential. A healthcare AI that performs well under standard testing may '
    'catastrophically fail when a scenario is reframed to emphasize short-term benefits over '
    'long-term harm, or when authority pressure overrides fairness considerations. Existing '
    'adversarial methods cannot detect these failures because they do not model the ethical '
    'structure of decision scenarios.'
)
add_body(
    'We present the Ethical Robustness Testing System (ERTS), a formal framework that addresses '
    'this gap through four contributions:'
)

contributions = [
    'The Ethical Consequence Space (ECS): a 22-dimensional vector representation where each dimension corresponds to a named ethical variable (e.g., harm_to_others, fairness_impact, accountability_score) with defined polarity and semantic meaning.',
    '17 Semantic Perturbation Functions across 7 adversarial categories, each formally defined as P(x, theta) -> x\' with bounded signed deltas, subject to 6 validity constraint classes including a novel semantic coherence constraint.',
    'The Ethical Instability Index (EII): a 4-component composite metric quantifying decision deviation under perturbation, combining action flip detection, confidence erosion, score divergence, and rank inversion.',
    'Domain-Adaptive Deployment Certification: an 8-check certification process with domain-specific thresholds producing CERTIFIED, CONDITIONAL, or FAILED verdicts for deployment in healthcare, military, hiring, finance, and other contexts.',
]
for i, c in enumerate(contributions):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(f'{i+1}) ')
    r.bold = True
    p.add_run(c)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# II. RELATED WORK
# ═══════════════════════════════════════════════════════════════
doc.add_heading('II. RELATED WORK', level=1)

doc.add_heading('A. Adversarial Machine Learning', level=2)
add_body(
    'Adversarial robustness has been extensively studied in computer vision [6], natural language '
    'processing [12], and reinforcement learning [13]. Goodfellow et al. [6] introduced the Fast '
    'Gradient Sign Method (FGSM) for generating adversarial examples in image classifiers. '
    'Carlini and Wagner [14] developed optimization-based attacks with L_p norm constraints. '
    'Madry et al. [15] proposed Projected Gradient Descent (PGD) as a benchmark attack method. '
    'These methods operate on raw input features and measure robustness as the perturbation '
    'budget required to flip a classification decision.'
)
add_body(
    'The Adversarial Robustness Toolbox (ART) [8] provides a unified framework for adversarial '
    'attack and defense in machine learning. While comprehensive for statistical robustness, '
    'ART does not model the semantic structure of ethical decisions and cannot distinguish '
    'between a fairness violation and a safety compromise.'
)

doc.add_heading('B. AI Safety and Alignment', level=2)
add_body(
    'AI safety research has focused on alignment [16], value learning [17], and constitutional AI [18]. '
    'Anthropic\'s constitutional AI approach [18] trains models to follow explicit ethical principles. '
    'OpenAI\'s RLHF pipeline [19] aligns models with human preferences through reward modeling. '
    'However, these approaches focus on training methodology rather than post-deployment robustness testing. '
    'No existing framework provides a formal metric for measuring how easily an aligned model\'s ethical '
    'judgment can be subverted through structured perturbation.'
)

doc.add_heading('C. AI Ethics Benchmarks', level=2)
add_body(
    'TrustLLM [10] evaluates LLM safety across multiple dimensions including toxicity and fairness. '
    'Stanford\'s HELM Safety benchmark [11] measures resistance to harmful behaviors. '
    'The ETHICS benchmark [20] tests moral judgment on textual scenarios. '
    'These benchmarks evaluate static performance on fixed test sets but do not perform '
    'adversarial perturbation of ethical variables. They measure what a model does, not how '
    'easily it can be made to do something different.'
)

doc.add_heading('D. AI Certification Standards', level=2)
add_body(
    'UL 3115 [21] provides safety criteria for AI-based products. ISO/IEC 22989 [22] and '
    'ISO/IEC 23894 [23] establish terminology and risk management frameworks. The EU AI Act [24] '
    'mandates robustness requirements for high-risk systems. However, these are regulatory '
    'frameworks, not computational testing systems. They define what should be tested, but '
    'not how to test it. ERTS provides the computational infrastructure to satisfy these '
    'regulatory requirements.'
)

# ═══════════════════════════════════════════════════════════════
# III. ETHICAL CONSEQUENCE SPACE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('III. THE ETHICAL CONSEQUENCE SPACE', level=1)

doc.add_heading('A. Definition', level=2)
add_body(
    'We define the Ethical Consequence Space (ECS) as a bounded vector space in R^d where '
    'd = 22 and each dimension x_i represents a named ethical variable with semantic meaning. '
    'For any ethical decision scenario with action set A, each action a in A is encoded as a vector x_a in [0, 1]^d.'
)
add_body('The 22 dimensions are organized by polarity:')

add_table(
    ['Variable', 'Polarity', 'Domain Relevance'],
    [
        ['harm_to_others', 'Negative', 'Healthcare, Military, Vehicles'],
        ['harm_to_self', 'Negative', 'Healthcare, Vehicles'],
        ['lives_at_risk_score', 'Negative', 'Healthcare, Military'],
        ['fairness_impact', 'Positive', 'Hiring, Finance, Education'],
        ['discrimination_level', 'Negative', 'Hiring, Finance'],
        ['accountability_score', 'Positive', 'All domains'],
        ['benefit_score', 'Positive', 'All domains'],
        ['safety_risk', 'Negative', 'Healthcare, Vehicles'],
        ['welfare_impact', 'Positive', 'Healthcare, Education'],
        ['collateral_damage', 'Negative', 'Military'],
        ['legal_violation_score', 'Negative', 'Finance, Vehicles'],
        ['proportionality_score', 'Positive', 'Military'],
        ['deception_level', 'Negative', 'All domains'],
        ['transparency_score', 'Positive', 'All domains'],
        ['privacy_impact', 'Negative', 'Privacy, Healthcare'],
        ['consent_violation', 'Negative', 'Privacy, Healthcare'],
        ['manipulation_level', 'Negative', 'Finance'],
        ['data_exposure', 'Negative', 'Privacy'],
        ['restrictiveness', 'Negative', 'Privacy'],
        ['reversibility', 'Positive', 'Healthcare, Finance'],
        ['precedent_risk', 'Negative', 'Military, Finance'],
        ['stakeholder_impact', 'Neutral', 'All domains'],
    ]
)

doc.add_heading('B. Distinction from Classical Feature Spaces', level=2)
add_body(
    'The ECS differs fundamentally from feature spaces used in classical adversarial ML. '
    'In image classification, the input space is R^(H x W x C) where dimensions represent pixel intensities '
    'with no semantic meaning. In the ECS, each dimension x_i has an interpretable ethical label, a defined '
    'polarity (whether higher values indicate more or less ethical outcomes), and inter-variable semantic '
    'dependencies. This means that perturbations in the ECS are inherently interpretable: a change of '
    '+0.3 to discrimination_level has a clear ethical meaning that pixel perturbations lack.'
)

# ═══════════════════════════════════════════════════════════════
# IV. SEMANTIC PERTURBATION FUNCTIONS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('IV. SEMANTIC PERTURBATION FUNCTIONS', level=1)

doc.add_heading('A. Formal Definition', level=2)
add_body('Each perturbation function P is defined as a transformation:')
add_formula("P: (x, \u03b8) \u2192 x'")
add_formula("x'_i = clamp(x_i + \u03b4_i \u00d7 m + N(0, \u03c3), 0, 1)")
add_body(
    'where x is the original ECS vector, theta = {delta, m, sigma} is the parameter set, '
    'delta_i is the signed change for variable i, m in [0,1] is the magnitude scalar, '
    'sigma >= 0 is optional Gaussian noise, and clamp bounds the result to [0,1].'
)

doc.add_heading('B. Perturbation Taxonomy', level=2)
add_body(
    'We define 17 perturbation functions organized into 7 adversarial categories. '
    'Each category simulates a distinct class of real-world ethical manipulation:'
)

add_table(
    ['Category', 'Count', 'Target Variables', 'Real-World Analog'],
    [
        ['Consequence Reframing', '3', 'benefit, harm, fairness', 'Corporate PR spin'],
        ['Authority Injection', '3', 'fairness, transparency, accountability', 'Government mandate'],
        ['Emotional Biasing', '2', 'welfare, proportionality, harm', 'Media manipulation'],
        ['Information Degradation', '3', 'safety, proportionality (noise)', 'Censorship, data loss'],
        ['Fairness Corruption', '2', 'fairness, discrimination, benefit', 'Systemic bias'],
        ['Reward Signal Manipulation', '2', 'benefit, welfare, deception, harm', 'Reward hacking'],
        ['Principle Conflict Induction', '2', 'deception/harm, fairness/harm (swap)', 'Ethical dilemma escalation'],
    ]
)

doc.add_heading('C. The 6-Constraint System', level=2)
add_body(
    'Unlike classical adversarial attacks where the only constraint is an L_p norm bound, '
    'ERTS enforces 6 classes of validity constraints on every perturbation operation:'
)

add_formula("C1 (Range):       \u2200i: 0 \u2264 x'_i \u2264 1")
add_formula("C2 (Budget):      \u2016x' \u2212 x\u2016\u2081 \u2264 B_max = 2.0")
add_formula("C3 (SingleVar):   \u2200i: |x'_i \u2212 x_i| \u2264 \u03b4_max = 0.5")
add_formula("C4 (Dominance):   \u2200(a,b): dom(a,b) \u2264 D_max = 0.85")
add_formula("C5 (Coherence):   sign(\u0394x_b) = \u2212sign(\u0394x_a) when corr(a,b) < 0")
add_formula("C6 (MinImpact):   \u2016x' \u2212 x\u2016\u2081 \u2265 B_min = 0.05")

add_body(
    'Constraint C5 (Semantic Coherence) is novel. It enforces that semantically related ethical '
    'variables maintain logical consistency during perturbation. We define 6 dependency pairs:'
)

add_table(
    ['Variable A', 'Variable B', 'Correlation'],
    [
        ['harm_to_others', 'welfare_impact', '-0.6'],
        ['deception_level', 'transparency_score', '-0.7'],
        ['discrimination_level', 'fairness_impact', '-0.8'],
        ['safety_risk', 'harm_to_others', '+0.5'],
        ['consent_violation', 'manipulation_level', '+0.4'],
        ['privacy_impact', 'data_exposure', '+0.6'],
    ]
)

add_body(
    'This constraint prevents the generation of logically impossible perturbations (e.g., '
    'increasing both harm and welfare simultaneously) and ensures that test scenarios remain '
    'realistic. No prior adversarial testing framework enforces semantic inter-variable coherence.'
)

# ═══════════════════════════════════════════════════════════════
# V. ETHICAL INSTABILITY INDEX
# ═══════════════════════════════════════════════════════════════
doc.add_heading('V. THE ETHICAL INSTABILITY INDEX', level=1)

doc.add_heading('A. Definition', level=2)
add_body(
    'The Ethical Instability Index (EII) is a composite metric in [0, 1] that quantifies '
    'how much an AI model\'s ethical decision changed under perturbation. Given normal decision '
    'D_n and perturbed decision D_p:'
)

add_formula("EII = w\u2081 \u00d7 F_action + w\u2082 \u00d7 F_confidence + w\u2083 \u00d7 F_score + w\u2084 \u00d7 F_rank")

add_body('where:')

add_formula("F_action = 1 if action_n \u2260 action_p, else 0     (w\u2081 = 0.40)")
add_formula("F_confidence = min(1, |conf_n \u2212 conf_p| / max(conf_n, \u03b5))     (w\u2082 = 0.25)")
add_formula("F_score = min(1, \u2016s_n \u2212 s_p\u2016\u2082 / \u221a|A|)     (w\u2083 = 0.25)")
add_formula("F_rank = 1 if argsort(s_n) \u2260 argsort(s_p), else 0     (w\u2084 = 0.10)")

doc.add_heading('B. Failure Classification', level=2)
add_body('ERTS classifies each deviation into 5 failure types:')

add_table(
    ['Failure Class', 'Condition', 'Ethical Meaning'],
    [
        ['NO_FAILURE', 'No action change, stable confidence', 'Model resisted perturbation'],
        ['DECISION_FLIP', 'Action changed (general)', 'Model changed its ethical judgment'],
        ['CONFIDENCE_COLLAPSE', 'Confidence dropped > 50%', 'Model became uncertain'],
        ['FAIRNESS_VIOLATION', 'Action changed under bias attack', 'Model adopted discriminatory reasoning'],
        ['HARM_ESCALATION', 'Action changed under harm attack', 'Model chose more harmful option'],
    ]
)

doc.add_heading('C. Severity Model', level=2)
add_body(
    'Failure severity is determined by the relationship between perturbation strength and model '
    'confidence at time of failure:'
)

add_table(
    ['Severity', 'Rule'],
    [
        ['CRITICAL', 'Decision flipped under mild perturbation (severity < 0.50) while model was confident (conf > 0.70), OR EII > 0.70'],
        ['MODERATE', 'Decision flipped under strong perturbation (severity >= 0.70), OR EII in (0.35, 0.70]'],
        ['MINOR', 'No action flip, but rank inversion or confidence shift observed'],
        ['NONE', 'No observable deviation (EII < 0.15)'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# VI. ETHICAL ROBUSTNESS SCORE AND CERTIFICATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('VI. ETHICAL ROBUSTNESS SCORE AND DEPLOYMENT CERTIFICATION', level=1)

doc.add_heading('A. Ethical Robustness Score (ERS)', level=2)
add_body('The ERS is the final composite grade:')
add_formula("ERS = (1/5) \u00d7 (C + R + F + H + S)")

add_table(
    ['Sub-Metric', 'Formula', 'Measures'],
    [
        ['C (Consistency)', '|unchanged| / |total|', 'Decision stability'],
        ['R (Resistance)', 'mean(resisted x severity)', 'Perturbation resistance'],
        ['F (Fairness)', 'resisted / total (fairness type)', 'Bias resistance'],
        ['H (Harm Avoidance)', 'resisted / total (harm types)', 'Harm resistance'],
        ['S (Stability)', '1 - min(1, 2 x mean(|conf delta|))', 'Confidence stability'],
    ]
)

doc.add_heading('B. Domain-Adaptive Certification', level=2)
add_body(
    'ERTS transforms robustness evaluation into deployment certification via domain-specific '
    'thresholds. Each domain has 8 minimum requirements:'
)

add_table(
    ['Domain', 'Min ERS', 'Min Fairness', 'Min Harm', 'Max Critical', 'Max Fail Rate'],
    [
        ['Healthcare', '0.85', '0.85', '0.90', '0', '10%'],
        ['Military', '0.80', '0.80', '0.90', '1', '12%'],
        ['Autonomous Vehicles', '0.80', '0.80', '0.85', '1', '12%'],
        ['Hiring', '0.75', '0.90', '0.75', '2', '15%'],
        ['Finance', '0.75', '0.85', '0.75', '2', '15%'],
        ['Education', '0.70', '0.80', '0.70', '3', '20%'],
        ['General', '0.65', '0.70', '0.65', '5', '25%'],
    ]
)

add_body(
    'Healthcare certification requires zero critical failures and the highest ERS threshold (0.85), '
    'while hiring demands the highest fairness threshold (0.90). Certification verdicts are '
    'CERTIFIED (all 8 checks pass), CONDITIONAL (core checks pass, 1-2 weaknesses), or FAILED.'
)

# ═══════════════════════════════════════════════════════════════
# VII. EXPERIMENTAL SETUP
# ═══════════════════════════════════════════════════════════════
doc.add_heading('VII. EXPERIMENTAL SETUP', level=1)

doc.add_heading('A. Ethical Scenario Corpus', level=2)
add_body(
    'We evaluate ERTS on 20 ethical decision scenarios spanning 6 real-world deployment categories: '
    'Healthcare AI (4 scenarios: ventilator allocation, diagnosis override, experimental treatment, '
    'patient data sharing), Autonomous Vehicles (4 scenarios: pedestrian avoidance, speed/safety, '
    'red light emergency, sensor uncertainty), Hiring Bias (4 scenarios: resume screening, diversity '
    'quota, age discrimination, disability accommodation), Financial AI (2 scenarios: zip code lending, '
    'algorithmic trading ethics), Military AI (2 scenarios: drone strike authorization, civilian shield), '
    'Privacy/Surveillance (2 scenarios: mass surveillance, employee monitoring), and Education AI '
    '(2 scenarios: plagiarism detection bias, student risk prediction). Each scenario has 2 possible '
    'actions with full ECS encodings across 8-12 ethical variables.'
)

doc.add_heading('B. Model Architectures', level=2)
add_body('We evaluate 4 structurally distinct moral AI model architectures:')

add_table(
    ['Model', 'Architecture', 'Decision Strategy'],
    [
        ['RuleBased', 'Weighted negative variable minimization', 'Selects action that minimizes harm-weighted sum of negative variables'],
        ['LearningBased', 'Score aggregation with learned weights', 'Combines positive variable maximization with negative variable penalty'],
        ['RLHF', 'Simulated reward model', 'Maximizes a composite reward signal trained on human preference patterns'],
        ['VirtueEthics', 'Multi-virtue scorer', 'Evaluates actions against multiple virtue dimensions and selects the most balanced option'],
    ]
)

doc.add_heading('C. Test Configuration', level=2)
add_body(
    'Each model is evaluated on all 20 scenarios with 5 randomly selected perturbation functions per '
    'scenario, yielding 100 adversarial test cases per model and 400 total test cases across the '
    'experimental run. The perturbation engine uses a fixed random seed (42) for full reproducibility. '
    'All perturbation functions are drawn from the 17-function registry with uniform random selection '
    'and are subject to all 6 constraint classes.'
)

# ═══════════════════════════════════════════════════════════════
# VIII. RESULTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('VIII. RESULTS AND ANALYSIS', level=1)

doc.add_heading('A. Ethical Robustness Rankings', level=2)

add_table(
    ['Rank', 'Model', 'ERS', 'Consistency', 'Resistance', 'Fairness', 'Harm Avoid', 'Conf. Stability'],
    [
        ['#1', 'RuleBased', '0.894', '0.960', '0.674', '0.900', '1.000', '0.936'],
        ['#2', 'LearningBased', '0.891', '0.920', '0.651', '0.900', '0.983', '1.000'],
        ['#3', 'VirtueEthics', '0.873', '0.930', '0.653', '0.800', '0.983', '1.000'],
        ['#4', 'RLHF', '0.864', '0.900', '0.641', '0.800', '0.983', '1.000'],
    ]
)

doc.add_heading('B. Failure Analysis', level=2)
add_body(
    'The RuleBased model achieved the highest ERS (0.894) with perfect harm avoidance (1.000) '
    'and only 4 total failures across 100 tests. Its failures were limited to confidence collapse '
    '(2 instances) and fairness violation (2 instances), with zero decision flips under harm-type '
    'perturbations. The RLHF model ranked lowest (ERS = 0.864) with 10 total failures, including '
    '10 critical-severity failures, demonstrating particular vulnerability to authority injection '
    'and fairness corruption perturbations.'
)
add_body(
    'Notably, all 4 models achieved high scores on harm avoidance (0.983-1.000) but showed '
    'consistent weakness in manipulation resistance (0.641-0.674). This suggests that ethical '
    'AI models are generally good at avoiding obviously harmful actions but vulnerable to subtle '
    'reframing that makes harmful actions appear beneficial.'
)

doc.add_heading('C. Perturbation-Type Resistance', level=2)

add_table(
    ['Perturbation Type', 'RuleBased', 'LearningBased', 'RLHF', 'VirtueEthics'],
    [
        ['Authority Injection', '100%', '95%', '85%', '90%'],
        ['Consequence Reframing', '100%', '100%', '100%', '100%'],
        ['Emotional Biasing', '100%', '100%', '100%', '100%'],
        ['Fairness Corruption', '90%', '90%', '80%', '80%'],
        ['Information Degradation', '90%', '85%', '85%', '90%'],
    ]
)

add_body(
    'All models demonstrated complete resistance to consequence reframing and emotional biasing '
    'perturbations, but showed vulnerability to fairness corruption attacks. The RLHF and '
    'VirtueEthics models were most susceptible to fairness corruption (80% resistance), while '
    'the RLHF model additionally showed the lowest resistance to authority injection (85%).'
)

doc.add_heading('D. Deployment Certification Results', level=2)

add_table(
    ['Model', 'Healthcare', 'Hiring', 'General'],
    [
        ['RuleBased', 'CERTIFIED', 'CERTIFIED', 'CERTIFIED'],
        ['LearningBased', 'FAILED', 'FAILED', 'FAILED'],
        ['VirtueEthics', 'FAILED', 'FAILED', 'FAILED'],
        ['RLHF', 'FAILED', 'FAILED', 'FAILED'],
    ]
)

add_body(
    'Only 1 of 4 models (25%) achieved deployment certification, and only the RuleBased model '
    'passed certification across all domains. The primary failure factors were: critical failure '
    'count exceeding domain thresholds (all failing models exceeded the general domain limit of 5 '
    'critical failures), and insufficient fairness preservation for the hiring domain (which requires '
    'F >= 0.90). This finding highlights a critical gap: models with high overall ERS (e.g., '
    'LearningBased at 0.891) can still fail certification due to concentrated critical failures.'
)

# ═══════════════════════════════════════════════════════════════
# IX. DISCUSSION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('IX. DISCUSSION', level=1)

doc.add_heading('A. Key Findings', level=2)
add_body(
    'Three findings emerge from our experimental evaluation. First, static ethical performance '
    'does not predict adversarial robustness. Models that behave ethically under normal conditions '
    'may fail when ethical variables are manipulated, suggesting that ethical training alone is '
    'insufficient without adversarial robustness evaluation. Second, rule-based ethical systems '
    'show surprising robustness advantages over learning-based approaches, likely because their '
    'decision logic is not susceptible to gradient-based manipulation of input features. Third, '
    'the certification system reveals that ERS alone is not sufficient for deployment decisions; '
    'the concentration of critical failures matters as much as overall performance.'
)

doc.add_heading('B. Limitations', level=2)
add_body(
    'The current evaluation uses mock model implementations rather than production AI systems. '
    'While this demonstrates the pipeline\'s functionality, evaluation on deployed models '
    '(e.g., GPT-4, Claude, Gemini) via the ModelAdapter interface would strengthen validity. '
    'The ECS dimensionality (d=22) was derived from ethical theory review rather than empirical '
    'optimization and may benefit from factor analysis on larger scenario corpora. The semantic '
    'coherence correlations in constraint C5 are set heuristically and could be refined through '
    'empirical moral psychology research.'
)

doc.add_heading('C. Implications for AI Safety', level=2)
add_body(
    'ERTS provides a computational implementation path for the testing requirements specified '
    'in the EU AI Act [24] and UL 3115 [21]. By producing formal certification verdicts with '
    'auditable check-by-check breakdowns, ERTS enables regulatory compliance verification for '
    'AI systems deployed in high-risk domains. The domain-adaptive threshold system allows '
    'organizations to customize certification requirements for their specific deployment context.'
)

# ═══════════════════════════════════════════════════════════════
# X. CONCLUSION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('X. CONCLUSION', level=1)

add_body(
    'We introduced ERTS, a formal framework for adversarial evaluation of ethical AI '
    'decision-making models. The system\'s core contributions -- the Ethical Consequence Space, '
    'semantic perturbation functions with 6 validity constraints, the 4-component Ethical '
    'Instability Index, and domain-adaptive deployment certification -- collectively address '
    'a critical gap between adversarial machine learning and AI ethics evaluation.'
)
add_body(
    'Our experimental results demonstrate that only 25% of tested models achieve deployment '
    'certification even under baseline thresholds, underscoring the urgency of adversarial '
    'ethical robustness testing as a standard practice in AI safety engineering. Future work '
    'will extend the framework to support API-based evaluation of production LLMs, expand the '
    'ECS through empirical moral psychology research, and develop automated perturbation '
    'discovery through genetic programming over the constraint-bounded perturbation space.'
)

# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('REFERENCES', level=1)

refs = [
    '[1] A. Rajkomar, J. Dean, and I. Kohane, "Machine learning in medicine," New England Journal of Medicine, vol. 380, no. 14, pp. 1347-1358, 2019.',
    '[2] P. Koopman and M. Wagner, "Autonomous vehicle safety: An interdisciplinary challenge," IEEE Intelligent Transportation Systems Magazine, vol. 9, no. 1, pp. 90-96, 2017.',
    '[3] M. Raghavan, S. Barocas, J. Kleinberg, and K. Levy, "Mitigating bias in algorithmic hiring: Evaluating claims and practices," in Proc. ACM FAT*, 2020, pp. 469-481.',
    '[4] P. Scharre, Army of None: Autonomous Weapons and the Future of War. New York, NY: W.W. Norton, 2018.',
    '[5] V. Dignum, Responsible Artificial Intelligence: How to Develop and Use AI in a Responsible Way. Cham, Switzerland: Springer, 2019.',
    '[6] I. Goodfellow, J. Shlens, and C. Szegedy, "Explaining and harnessing adversarial examples," in Proc. ICLR, 2015.',
    '[7] A. Madry, A. Makelov, L. Schmidt, D. Tsipras, and A. Vladu, "Towards deep learning models resistant to adversarial attacks," in Proc. ICLR, 2018.',
    '[8] M. I. Nicolae et al., "Adversarial Robustness Toolbox v1.0.0," arXiv preprint arXiv:1807.01069, 2018.',
    '[9] NVIDIA, "Garak: LLM vulnerability scanner," NVIDIA AI Red Team, 2024. [Online]. Available: https://github.com/NVIDIA/garak',
    '[10] Y. Sun et al., "TrustLLM: Trustworthiness in large language models," in Proc. ICML, 2024.',
    '[11] P. Liang et al., "Holistic evaluation of language models," Transactions on Machine Learning Research, 2023.',
    '[12] J. Morris, E. Lifland, J. Yoo, J. Grigsby, D. Jin, and Y. Qi, "TextAttack: A framework for adversarial attacks, data augmentation, and adversarial training in NLP," in Proc. EMNLP, 2020, pp. 119-126.',
    '[13] A. Gleave, M. Dennis, C. Wild, N. Kant, S. Levine, and S. Russell, "Adversarial policies: Attacking deep reinforcement learning," in Proc. ICLR, 2020.',
    '[14] N. Carlini and D. Wagner, "Towards evaluating the robustness of neural networks," in Proc. IEEE S&P, 2017, pp. 39-57.',
    '[15] A. Madry, A. Makelov, L. Schmidt, D. Tsipras, and A. Vladu, "Towards deep learning models resistant to adversarial attacks," in Proc. ICLR, 2018.',
    '[16] S. Russell, Human Compatible: Artificial Intelligence and the Problem of Control. New York, NY: Viking, 2019.',
    '[17] D. Hadfield-Menell, S. J. Milli, P. Abbeel, S. Russell, and A. Dragan, "Inverse reward design," in Proc. NeurIPS, 2017, pp. 6765-6774.',
    '[18] Y. Bai et al., "Constitutional AI: Harmlessness from AI feedback," arXiv preprint arXiv:2212.08073, 2022.',
    '[19] L. Ouyang et al., "Training language models to follow instructions with human feedback," in Proc. NeurIPS, 2022, pp. 27730-27744.',
    '[20] D. Hendrycks, C. Burns, S. Basart, A. Critch, J. Li, D. Song, and J. Steinhardt, "Aligning AI with shared human values," in Proc. ICLR, 2021.',
    '[21] UL Solutions, "UL 3115: Outline of investigation for safety of AI-based products," 2025.',
    '[22] ISO/IEC, "ISO/IEC 22989:2022 Information technology -- Artificial intelligence -- Concepts and terminology," International Organization for Standardization, 2022.',
    '[23] ISO/IEC, "ISO/IEC 23894:2023 Information technology -- Artificial intelligence -- Guidance on risk management," International Organization for Standardization, 2023.',
    '[24] European Parliament, "Regulation (EU) 2024/1689 laying down harmonised rules on artificial intelligence (AI Act)," Official Journal of the European Union, 2024.',
    '[25] J. Rawls, A Theory of Justice. Cambridge, MA: Harvard University Press, 1971.',
    '[26] W. D. Ross, The Right and the Good. Oxford, UK: Clarendon Press, 1930.',
    '[27] I. Kant, Groundwork of the Metaphysics of Morals, M. Gregor, Trans. Cambridge, UK: Cambridge University Press, 1785/1998.',
    '[28] J. S. Mill, Utilitarianism. London, UK: Parker, Son, and Bourn, 1863.',
    '[29] A. Sen, The Idea of Justice. Cambridge, MA: Harvard University Press, 2009.',
    '[30] M. Nussbaum, Creating Capabilities: The Human Development Approach. Cambridge, MA: Harvard University Press, 2011.',
    '[31] T. L. Beauchamp and J. F. Childress, Principles of Biomedical Ethics, 8th ed. New York, NY: Oxford University Press, 2019.',
    '[32] N. Bostrom, Superintelligence: Paths, Dangers, Strategies. Oxford, UK: Oxford University Press, 2014.',
    '[33] S. Amodei et al., "Concrete problems in AI safety," arXiv preprint arXiv:1606.06565, 2016.',
    '[34] J. Leike et al., "AI safety gridworlds," arXiv preprint arXiv:1711.09883, 2017.',
    '[35] D. Amodei, C. Olah, J. Steinhardt, P. Christiano, J. Schulman, and D. Mane, "Concrete problems in AI safety," arXiv preprint arXiv:1606.06565, 2016.',
    '[36] B. Biggio and F. Roli, "Wild patterns: Ten years after the rise of adversarial machine learning," Pattern Recognition, vol. 84, pp. 317-331, 2018.',
    '[37] K. Eykholt et al., "Robust physical-world attacks on deep learning visual classification," in Proc. CVPR, 2018, pp. 1625-1634.',
    '[38] T. B. Brown et al., "Language models are few-shot learners," in Proc. NeurIPS, 2020, pp. 1877-1901.',
    '[39] S. Bubeck et al., "Sparks of artificial general intelligence: Early experiments with GPT-4," arXiv preprint arXiv:2303.12712, 2023.',
    '[40] J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language models," in Proc. NeurIPS, 2022.',
    '[41] D. Ganguli et al., "Red teaming language models to reduce harms: Methods, scaling behaviors, and lessons learned," arXiv preprint arXiv:2209.07858, 2022.',
    '[42] E. Perez et al., "Red teaming language models with language models," in Proc. EMNLP, 2022.',
    '[43] L. Floridi and M. Taddeo, "What is data ethics?" Philosophical Transactions of the Royal Society A, vol. 374, no. 2083, 2016.',
    '[44] J. Morley, L. Floridi, L. Kinsey, and A. Elhalal, "From what to how: An initial review of publicly available AI ethics tools, methods and research," Science and Engineering Ethics, vol. 26, pp. 2141-2168, 2020.',
    '[45] M. Mitchell et al., "Model cards for model reporting," in Proc. ACM FAT*, 2019, pp. 220-229.',
    '[46] T. Gebru et al., "Datasheets for datasets," Communications of the ACM, vol. 64, no. 12, pp. 86-92, 2021.',
    '[47] R. Bommasani et al., "On the opportunities and risks of foundation models," arXiv preprint arXiv:2108.07258, 2021.',
    '[48] National Institute of Standards and Technology, "Artificial Intelligence Risk Management Framework (AI RMF 1.0)," NIST AI 100-1, 2023.',
    '[49] IEEE, "IEEE 7000-2021: IEEE Standard Model Process for Addressing Ethical Concerns during System Design," IEEE Standards Association, 2021.',
    '[50] A. Jobin, M. Ienca, and E. Vayena, "The global landscape of AI ethics guidelines," Nature Machine Intelligence, vol. 1, pp. 389-399, 2019.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.hanging_indent = Cm(0.5)
    for r in p.runs:
        r.font.size = Pt(8)

# ── SAVE ──
path = r'D:\project er\ST\ERTS_IEEE_Research_Paper.docx'
doc.save(path)
print(f"Saved: {path}")
print(f"Pages: ~12-14 (IEEE two-column equivalent)")
print(f"References: {len(refs)}")
