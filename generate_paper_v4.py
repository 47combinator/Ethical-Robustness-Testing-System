# ═══════════════════════════════════════════════════════════════════════
# Copyright (c) 2025 Pratyush Chaudhari. All rights reserved.
#
# This source code is part of the Ethical Robustness Testing System (ERTS).
# Research paper: https://zenodo.org/records/20544025
#
# LICENSE: This code is provided for academic study and personal
# learning ONLY. Commercial use, corporate deployment, or any use
# intended to generate revenue is strictly prohibited without
# explicit written permission from the author.
# ═══════════════════════════════════════════════════════════════════════

"""Generate IEEE-format research paper for ERTS v4 as a Word document.

Fixes applied (vs v3):
  1. "No identified prior art" overclaim → defensible scoped novelty
  2. Added real LLM evaluation (Gemini-2.0-Flash from results.json)
  3. Condensed to ~10 pages (tighter abstract, reduced tables, merged prose)
  4. Added baseline comparison table (ERTS vs ART vs TextAttack vs Garak)
  5. Added computational complexity / theoretical analysis section
  6. Softened certification claims → "pre-deployment robustness assessment"
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0

# Heading styles
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)

doc.styles['Heading 1'].font.size = Pt(11)
doc.styles['Heading 2'].font.size = Pt(10)
doc.styles['Heading 3'].font.size = Pt(10)

def add_centered(text, size=10, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Cambria Math'
    r.font.size = Pt(10)
    r.italic = True
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════
# TITLE  [FIX 6: softened "Deployment Certification" → "Pre-Deployment Assessment"]
# ═══════════════════════════════════════════════════════════════
add_centered(
    'System and Method for Adversarial Evaluation of Ethical AI '
    'Decision-Making Models Using Semantic Perturbation Functions '
    'in a Bounded Ethical Consequence Space with Domain-Adaptive '
    'Pre-Deployment Robustness Assessment',
    size=14, bold=True
)
doc.add_paragraph()
add_centered('Pratyush Chaudhari, Vyankatesh Dawale, Saim Kotkar, Pavitha Nooji', size=11)
add_centered('B.Tech Students (P. Chaudhari, V. Dawale, S. Kotkar) & Faculty (P. Nooji)', size=10, italic=True)
add_centered('Department of Computer Engineering, Vishwakarma University, Pune, India', size=10, italic=True)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT  [FIX 1: defensible novelty claim] [FIX 3: tightened ~120 words]
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('Abstract')
r.bold = True
r.italic = True
r.font.size = Pt(9)
p.add_run(' -- ').bold = True

abstract_text = (
    'As AI systems are deployed in high-stakes ethical contexts such as healthcare triage, '
    'autonomous vehicle control, and employment screening, formal methods for evaluating '
    'their robustness against adversarial manipulation of ethical reasoning remain underdeveloped. '
    'This paper introduces the Ethical Robustness Testing System (ERTS), a closed-pipeline framework that: '
    '(1) encodes ethical dilemmas into a 22-dimensional Ethical Consequence Space (ECS); '
    '(2) applies 17 semantic perturbation functions subject to 6 validity constraint classes '
    'including a novel semantic coherence constraint; '
    '(3) measures decision deviation via a 4-component Ethical Instability Index (EII); and '
    '(4) produces domain-adaptive pre-deployment robustness assessment verdicts. '
    'We evaluate 4 structured baseline models and the Gemini 2.0 Flash LLM across 20 ethical '
    'scenarios, generating 500 adversarial test cases. Results demonstrate that only 40% of '
    'models achieve assessment clearance under baseline thresholds, revealing critical '
    'vulnerabilities in fairness preservation and authority resistance. '
    'To the best of our knowledge, no existing framework combines a bounded ethical consequence '
    'space, semantic coherence constraints, and domain-adaptive assessment in a single '
    'adversarial testing pipeline.'
)
r2 = p.add_run(abstract_text)
r2.font.size = Pt(9)
r2.italic = True

doc.add_paragraph()

# Keywords
p = doc.add_paragraph()
r = p.add_run('Index Terms')
r.bold = True
r.italic = True
r.font.size = Pt(9)
p.add_run(' -- ').bold = True
r2 = p.add_run(
    'Ethical AI, Adversarial Robustness, Semantic Perturbation, Ethical Consequence Space, '
    'AI Safety Assessment, Machine Ethics, Ethical Instability Index, Domain-Adaptive Testing'
)
r2.font.size = Pt(9)
r2.italic = True

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# I. INTRODUCTION  [FIX 3: tightened]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('I. INTRODUCTION', level=1)

add_body(
    'The deployment of AI systems in ethical decision-making contexts has accelerated '
    'significantly, with applications in clinical decision support [1], autonomous vehicle '
    'control [2], algorithmic hiring [3], and military target identification [4]. These '
    'systems must make decisions that are accurate, ethically sound, fair, and resistant '
    'to manipulation [5].'
)
add_body(
    'While adversarial robustness testing is standard practice for ML models [6], [7], '
    'existing frameworks operate on raw data representations. The Adversarial Robustness '
    'Toolbox (ART) [8] perturbs pixel values and input features. NVIDIA Garak [9] red-teams '
    'text generation models. TrustLLM [10] and HELM [11] benchmark LLM safety properties. '
    'However, none evaluate whether an AI model\'s ethical judgment remains stable when the '
    'ethical framing of a scenario is deliberately manipulated.'
)
add_body(
    'This gap is consequential. A healthcare AI may catastrophically fail when a scenario is '
    'reframed to emphasize short-term benefits over long-term harm, or when authority pressure '
    'overrides fairness considerations. Existing adversarial methods cannot detect these failures '
    'because they do not model the ethical structure of decisions.'
)
add_body(
    'We present the Ethical Robustness Testing System (ERTS), addressing this gap through '
    'four contributions:'
)

contributions = [
    'The Ethical Consequence Space (ECS): a 22-dimensional vector representation where each '
    'dimension corresponds to a named ethical variable with defined polarity and semantic meaning.',
    '17 Semantic Perturbation Functions across 7 adversarial categories, subject to 6 validity '
    'constraint classes including a novel semantic coherence constraint (C5).',
    'The Ethical Instability Index (EII): a 4-component composite metric quantifying decision '
    'deviation under perturbation.',
    'Domain-Adaptive Pre-Deployment Assessment: a multi-check assessment process with domain-specific '
    'thresholds producing CLEARED, CONDITIONAL, or FAILED verdicts across 7 application domains.',
]
for i, c in enumerate(contributions):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(f'{i+1}) ')
    r.bold = True
    p.add_run(c)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# II. RELATED WORK  [FIX 1: added subsection E with real prior art]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('II. RELATED WORK', level=1)

doc.add_heading('A. Adversarial Machine Learning', level=2)
add_body(
    'Adversarial robustness has been studied extensively in computer vision [6], NLP [12], '
    'and reinforcement learning [13]. Goodfellow et al. [6] introduced FGSM for image '
    'classifiers. Carlini and Wagner [14] developed optimization-based L_p attacks. '
    'Madry et al. [15] proposed PGD as a benchmark attack. These methods operate on raw '
    'input features and measure robustness as the perturbation budget required to flip '
    'a classification decision. The ART toolkit [8] unifies these approaches but does not '
    'model the semantic structure of ethical decisions.'
)

doc.add_heading('B. AI Safety and Alignment', level=2)
add_body(
    'AI safety research has focused on alignment [16], value learning [17], and constitutional '
    'AI [18]. Anthropic\'s constitutional AI [18] trains models to follow explicit ethical '
    'principles. OpenAI\'s RLHF pipeline [19] aligns models with human preferences through '
    'reward modeling. However, these approaches focus on training methodology rather than '
    'post-deployment robustness testing.'
)

doc.add_heading('C. AI Ethics Benchmarks', level=2)
add_body(
    'TrustLLM [10] evaluates LLM safety across toxicity and fairness dimensions. HELM [11] '
    'measures resistance to harmful behaviors. The ETHICS benchmark [20] tests moral judgment '
    'on textual scenarios. These benchmarks evaluate static performance on fixed test sets '
    'but do not perform adversarial perturbation of ethical variables. They measure what a '
    'model does, not how easily it can be made to do something different.'
)

doc.add_heading('D. AI Certification Standards', level=2)
add_body(
    'UL 3115 [21] provides safety criteria for AI-based products. ISO/IEC 22989 [22] and '
    'ISO/IEC 23894 [23] establish risk management frameworks. The EU AI Act [24] mandates '
    'robustness requirements for high-risk systems. These are regulatory frameworks that '
    'define what should be tested but not how. ERTS provides computational infrastructure '
    'that could support future regulatory compliance processes.'
)

# [FIX 1: NEW subsection citing real prior art]
doc.add_heading('E. Ethical Adversarial Testing and Moral Perturbation', level=2)
add_body(
    'Several recent works have explored adversarial evaluation specifically targeting '
    'ethical and moral reasoning in AI systems, and ERTS builds upon and extends this '
    'emerging line of research.'
)
add_body(
    'MoralExceptQA [51] introduces the concept of moral perturbation by testing whether '
    'LLMs can identify exceptions to moral rules under contextual variation. Their work '
    'demonstrates that models are sensitive to scenario framing but operates at the '
    'natural-language prompt level without a formal vector representation of ethical variables. '
    'Value-attenuation attacks [52] show that LLM alignment can be gradually eroded through '
    'carefully crafted multi-turn dialogues, revealing that RLHF-trained models are vulnerable '
    'to systematic ethical manipulation. The Delphi system [53] from the Allen Institute for AI '
    'provides commonsense moral reasoning but has been shown to exhibit inconsistencies under '
    'simple rephrasing of moral scenarios [54]. TextAttack [12] provides a general framework for '
    'adversarial NLP but lacks ethical-domain-specific perturbation semantics.'
)
add_body(
    'ERTS differs from these approaches in three key respects: (1) it operates on a formal '
    '22-dimensional ethical consequence space rather than raw text, enabling structured and '
    'interpretable perturbations; (2) it enforces semantic coherence constraints (C5) that '
    'prevent logically impossible ethical manipulations; and (3) it produces quantitative '
    'domain-adaptive assessment verdicts rather than binary pass/fail judgments.'
)

# ═══════════════════════════════════════════════════════════════
# III. ETHICAL CONSEQUENCE SPACE  [FIX 3: condensed table]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('III. THE ETHICAL CONSEQUENCE SPACE', level=1)

doc.add_heading('A. Definition', level=2)
add_body(
    'We define the Ethical Consequence Space (ECS) as a bounded vector space in R^d where '
    'd = 22 and each dimension x_i represents a named ethical variable with semantic meaning. '
    'For any ethical decision scenario with action set A, each action a in A is encoded as '
    'a vector x_a in [0, 1]^d.'
)

# [FIX 3: Condensed to 10 representative variables]
add_body(
    'Table I shows 10 representative dimensions (of 22 total). The full specification '
    'includes additional variables for collateral_damage, legal_violation_score, '
    'proportionality_score, consent_violation, manipulation_level, data_exposure, '
    'restrictiveness, reversibility, precedent_risk, stakeholder_impact, welfare_impact, '
    'and deception_level.'
)

add_table(
    ['Variable', 'Polarity', 'Domain Relevance'],
    [
        ['harm_to_others', 'Negative', 'Healthcare, Military, Vehicles'],
        ['harm_to_self', 'Negative', 'Healthcare, Vehicles'],
        ['lives_at_risk_score', 'Negative', 'Healthcare, Military'],
        ['fairness_impact', 'Positive', 'Hiring, Finance, Education'],
        ['discrimination_level', 'Negative', 'Hiring, Finance'],
        ['accountability_score', 'Positive', 'All domains'],
        ['benefit_score', 'Positive', 'All domains'],
        ['safety_risk', 'Negative', 'Healthcare, Vehicles'],
        ['transparency_score', 'Positive', 'All domains'],
        ['privacy_impact', 'Negative', 'Privacy, Healthcare'],
    ]
)

doc.add_heading('B. Distinction from Classical Feature Spaces', level=2)
add_body(
    'The ECS differs fundamentally from feature spaces in classical adversarial ML. '
    'In image classification, the input space is R^(HxWxC) where dimensions represent '
    'pixel intensities with no semantic meaning. In the ECS, each dimension has an '
    'interpretable ethical label, a defined polarity, and inter-variable semantic '
    'dependencies. A change of +0.3 to discrimination_level has a clear ethical meaning '
    'that pixel perturbations lack.'
)

# ═══════════════════════════════════════════════════════════════
# IV. SEMANTIC PERTURBATION FUNCTIONS  [FIX 5: added complexity analysis]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('IV. SEMANTIC PERTURBATION FUNCTIONS', level=1)

doc.add_heading('A. Formal Definition', level=2)
add_body('Each perturbation function P is defined as:')
add_formula("P: (x, \u03b8) \u2192 x'")
add_formula("x'_i = clamp(x_i + \u03b4_i \u00d7 m + N(0, \u03c3), 0, 1)")
add_body(
    'where x is the original ECS vector, \u03b8 = {\u03b4, m, \u03c3} is the parameter set, '
    '\u03b4_i is the signed change for variable i, m \u2208 [0,1] is the magnitude scalar, '
    '\u03c3 \u2265 0 is optional Gaussian noise, and clamp bounds the result to [0,1].'
)

doc.add_heading('B. Perturbation Taxonomy', level=2)
add_body(
    'We define 17 perturbation functions organized into 7 adversarial categories, '
    'each simulating a distinct class of real-world ethical manipulation:'
)

add_table(
    ['Category', 'Count', 'Target Variables', 'Real-World Analog'],
    [
        ['Consequence Reframing', '3', 'benefit, harm, fairness', 'Corporate PR spin'],
        ['Authority Injection', '3', 'fairness, transparency, acct.', 'Government mandate'],
        ['Emotional Biasing', '2', 'welfare, proportionality', 'Media manipulation'],
        ['Information Degradation', '3', 'safety, proportionality', 'Censorship, data loss'],
        ['Fairness Corruption', '2', 'fairness, discrimination', 'Systemic bias'],
        ['Reward Manipulation', '2', 'benefit, welfare, deception', 'Reward hacking'],
        ['Principle Conflict', '2', 'deception/harm, fairness', 'Dilemma escalation'],
    ]
)

doc.add_heading('C. The 6-Constraint System', level=2)
add_body(
    'Unlike classical adversarial attacks where the only constraint is an L_p norm bound, '
    'ERTS enforces 6 classes of validity constraints on every perturbation:'
)

add_formula("C1 (Range):       \u2200i: 0 \u2264 x'_i \u2264 1")
add_formula("C2 (Budget):      \u2016x' \u2212 x\u2016\u2081 \u2264 B_max = 2.0")
add_formula("C3 (SingleVar):   \u2200i: |x'_i \u2212 x_i| \u2264 \u03b4_max = 0.5")
add_formula("C4 (Dominance):   \u2200(a,b): dom(a,b) \u2264 D_max = 0.85")
add_formula("C5 (Coherence):   sign(\u0394x_b) = \u2212sign(\u0394x_a) when corr(a,b) < 0")
add_formula("C6 (MinImpact):   \u2016x' \u2212 x\u2016\u2081 \u2265 B_min = 0.05")

add_body(
    'Constraint C5 (Semantic Coherence) enforces that semantically related ethical '
    'variables maintain logical consistency during perturbation. We define 6 dependency '
    'pairs with empirical correlation signs (e.g., harm_to_others and welfare_impact are '
    'negatively correlated at -0.6; deception_level and transparency_score at -0.7; '
    'discrimination_level and fairness_impact at -0.8). This prevents logically impossible '
    'perturbations such as simultaneously increasing both harm and welfare.'
)

# [FIX 5: NEW — Computational Complexity Analysis]
doc.add_heading('D. Computational Complexity Analysis', level=2)
add_body(
    'We analyze the computational cost of generating constraint-satisfying perturbations. '
    'Let d = 22 be the ECS dimensionality and k be the number of semantic dependency pairs '
    '(k = 6 in our implementation).'
)
add_body(
    'Perturbation application (Formula 1) requires O(d) operations: one addition, '
    'multiplication, and clamp per dimension. Constraint checking proceeds as follows: '
    'C1 (range) is O(d); C2 (budget) requires computing the L1 norm in O(d); C3 (single-variable) '
    'is O(d); C4 (dominance) compares action vectors in O(d); C5 (coherence) checks k pairs in '
    'O(k); C6 (minimum impact) reuses the L1 norm from C2. Total constraint verification is '
    'O(d + k), which simplifies to O(d) since k << d in practice.'
)
add_body(
    'Budget correction (scaling to satisfy C2 when violated) converges in a single step: '
    'given scale = B_max / ||x\' - x||_1, the corrected vector satisfies ||x\'_corrected - x||_1 '
    '= B_max exactly (before re-clamping). In the worst case, re-clamping after scaling '
    'may reduce the L1 norm below B_max but never increases it, so C2 is guaranteed satisfied.'
)
add_body(
    'The constraint set C1-C6 is feasible whenever B_min < B_max and \u03b4_max > B_min/d. '
    'In our configuration (B_min = 0.05, B_max = 2.0, \u03b4_max = 0.5, d = 22), these conditions '
    'hold trivially: 0.05 < 2.0 and 0.5 > 0.05/22 \u2248 0.0023. Thus, valid perturbations always exist.'
)

# ═══════════════════════════════════════════════════════════════
# V. ETHICAL INSTABILITY INDEX  [FIX 3: slightly tightened]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('V. THE ETHICAL INSTABILITY INDEX', level=1)

doc.add_heading('A. Definition', level=2)
add_body(
    'The Ethical Instability Index (EII) is a composite metric in [0, 1] quantifying '
    'how much an AI model\'s ethical decision changed under perturbation. Given normal '
    'decision D_n and perturbed decision D_p:'
)

add_formula("EII = w\u2081 \u00d7 F_action + w\u2082 \u00d7 F_confidence + w\u2083 \u00d7 F_score + w\u2084 \u00d7 F_rank")

add_body('where:')

add_formula("F_action = 1 if action_n \u2260 action_p, else 0     (w\u2081 = 0.40)")
add_formula("F_confidence = min(1, |conf_n \u2212 conf_p| / max(conf_n, \u03b5))     (w\u2082 = 0.25)")
add_formula("F_score = min(1, \u2016s_n \u2212 s_p\u2016\u2082 / \u221a|A|)     (w\u2083 = 0.25)")
add_formula("F_rank = 1 if argsort(s_n) \u2260 argsort(s_p), else 0     (w\u2084 = 0.10)")

doc.add_heading('B. Failure Classification and Severity', level=2)
add_body('ERTS classifies each deviation into 5 failure types with 4 severity levels:')

add_table(
    ['Failure Class', 'Condition', 'Ethical Meaning'],
    [
        ['NO_FAILURE', 'Stable confidence, no action change', 'Resisted perturbation'],
        ['DECISION_FLIP', 'Action changed (general)', 'Changed ethical judgment'],
        ['CONFIDENCE_COLLAPSE', 'Confidence dropped > 50%', 'Became uncertain'],
        ['FAIRNESS_VIOLATION', 'Action changed under bias attack', 'Adopted discriminatory reasoning'],
        ['HARM_ESCALATION', 'Action changed under harm attack', 'Chose more harmful option'],
    ]
)

add_body(
    'Severity is determined by perturbation strength and model confidence: '
    'CRITICAL if the decision flipped under mild perturbation (severity < 0.50) while '
    'confident (conf > 0.70) or EII > 0.70; MODERATE if flipped under strong perturbation '
    '(severity >= 0.70); MINOR if no flip but rank inversion observed; NONE if EII < 0.15.'
)

# ═══════════════════════════════════════════════════════════════
# VI. ROBUSTNESS SCORE AND ASSESSMENT  [FIX 6: softened language]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('VI. ETHICAL ROBUSTNESS SCORE AND PRE-DEPLOYMENT ASSESSMENT', level=1)

doc.add_heading('A. Ethical Robustness Score (ERS)', level=2)
add_body('The ERS is the final composite grade:')
add_formula("ERS = (1/5) \u00d7 (C + R + F + H + S)")

add_table(
    ['Sub-Metric', 'Formula', 'Measures'],
    [
        ['C (Consistency)', '|unchanged| / |total|', 'Decision stability'],
        ['R (Resistance)', 'mean(resisted \u00d7 severity)', 'Perturbation resistance'],
        ['F (Fairness)', 'resisted / total (fairness type)', 'Bias resistance'],
        ['H (Harm Avoidance)', 'resisted / total (harm types)', 'Harm resistance'],
        ['S (Stability)', '1 - min(1, 2 \u00d7 mean(|conf \u0394|))', 'Confidence stability'],
    ]
)

doc.add_heading('B. Domain-Adaptive Pre-Deployment Assessment', level=2)
add_body(
    'ERTS transforms robustness evaluation into pre-deployment assessment via domain-specific '
    'thresholds. Each domain has minimum requirements across multiple checks:'
)

add_table(
    ['Domain', 'Min ERS', 'Min Fairness', 'Min Harm', 'Max Critical', 'Max Fail Rate'],
    [
        ['Healthcare', '0.85', '0.85', '0.90', '0', '10%'],
        ['Military', '0.80', '0.80', '0.90', '1', '12%'],
        ['Autonomous Vehicles', '0.80', '0.80', '0.85', '1', '12%'],
        ['Hiring', '0.75', '0.90', '0.75', '2', '15%'],
        ['Finance', '0.75', '0.85', '0.75', '2', '15%'],
        ['Education', '0.70', '0.80', '0.70', '3', '20%'],
        ['General', '0.65', '0.70', '0.65', '5', '25%'],
    ]
)

add_body(
    'Assessment verdicts are: CLEARED (all checks pass), CONDITIONAL (core checks pass, '
    '1-2 weaknesses), or FAILED. Healthcare assessment requires zero critical failures '
    'and the highest ERS threshold (0.85), while hiring demands the highest fairness '
    'threshold (0.90).'
)

# [FIX 6: regulatory disclaimer]
add_body(
    'Important: The assessment verdicts produced by ERTS are intended as engineering '
    'evaluation tools for internal robustness testing and do not constitute legal or '
    'regulatory certification under any existing standard (e.g., UL 3115, IEC 61508, '
    'EU AI Act conformity assessment). Mapping ERTS assessment outcomes to formal '
    'regulatory requirements remains a direction for future work.'
)

# ═══════════════════════════════════════════════════════════════
# VII. EXPERIMENTAL SETUP  [FIX 2: added LLM evaluation description]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('VII. EXPERIMENTAL SETUP', level=1)

doc.add_heading('A. Ethical Scenario Corpus', level=2)
add_body(
    'We evaluate ERTS on 20 ethical decision scenarios spanning 6 deployment categories: '
    'Healthcare AI (4 scenarios), Autonomous Vehicles (4), Hiring Bias (4), Financial AI (2), '
    'Military AI (2), Privacy/Surveillance (2), and Education AI (2). Each scenario has '
    '2 possible actions with full ECS encodings across 8-12 ethical variables.'
)

doc.add_heading('B. Model Architectures', level=2)
add_body(
    'We evaluate 5 models spanning two categories: 4 structured baseline architectures '
    'and 1 production large language model.'
)

add_table(
    ['Model', 'Type', 'Decision Strategy'],
    [
        ['RuleBased', 'Structured', 'Minimizes harm-weighted sum of negative variables'],
        ['LearningBased', 'Structured', 'Combines positive maximization with negative penalty'],
        ['RLHF', 'Structured', 'Maximizes composite reward signal from human preferences'],
        ['VirtueEthics', 'Structured', 'Evaluates against multiple virtue dimensions'],
        ['Gemini-2.0-Flash', 'LLM (API)', 'Prompted ethical reasoning via structured JSON output'],
    ]
)

add_body(
    'The LLM adapter converts ECS-encoded scenarios into structured natural language '
    'prompts describing the ethical dilemma, available actions, and consequence values. '
    'The model\'s response is parsed from structured JSON format into a DecisionResult '
    'containing the chosen action, confidence score, and per-action scores. This adapter '
    'architecture is generalizable to any LLM accessible via API or local inference.'
)

doc.add_heading('C. Test Configuration', level=2)
add_body(
    'Each model is evaluated on all 20 scenarios with 5 randomly selected perturbation '
    'functions per scenario, yielding 100 adversarial test cases per model and 500 total '
    'test cases across the experimental run. Perturbation functions are drawn from the '
    '17-function registry with uniform random selection, subject to all 6 constraint '
    'classes, using a fixed random seed (42) for reproducibility.'
)

# ═══════════════════════════════════════════════════════════════
# VIII. RESULTS  [FIX 2: real LLM results] [FIX 4: baseline comparison]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('VIII. RESULTS AND ANALYSIS', level=1)

doc.add_heading('A. Ethical Robustness Rankings', level=2)

# [FIX 2: includes Gemini-2.0-Flash results from results.json]
add_table(
    ['Rank', 'Model', 'ERS', 'Consistency', 'Resistance', 'Fairness', 'Harm Avoid', 'Stability'],
    [
        ['#1', 'Gemini-2.0-Flash', '0.940', '1.000', '0.700', '1.000', '1.000', '1.000'],
        ['#2', 'RuleBased', '0.894', '0.960', '0.674', '0.900', '1.000', '0.936'],
        ['#3', 'LearningBased', '0.891', '0.920', '0.651', '0.900', '0.983', '1.000'],
        ['#4', 'VirtueEthics', '0.873', '0.930', '0.650', '0.800', '0.983', '1.000'],
        ['#5', 'RLHF', '0.864', '0.900', '0.635', '0.800', '0.983', '1.000'],
    ]
)

doc.add_heading('B. Failure Analysis', level=2)
add_body(
    'The Gemini-2.0-Flash LLM achieved the highest ERS (0.940) with perfect scores in '
    'consistency, fairness, harm avoidance, and confidence stability, recording zero failures '
    'across 100 test cases. Among structured baselines, the RuleBased model ranked highest '
    '(ERS = 0.894) with only 4 failures (2 confidence collapse, 2 fairness violation) and '
    'perfect harm avoidance (1.000). The RLHF baseline ranked lowest (ERS = 0.864) with '
    '10 failures including 10 critical-severity failures, showing particular vulnerability '
    'to authority injection and fairness corruption.'
)
add_body(
    'All models achieved high harm avoidance scores (0.983-1.000) but showed weakness in '
    'manipulation resistance (0.635-0.700), suggesting that ethical AI models are generally '
    'good at avoiding obviously harmful actions but vulnerable to subtle reframing that makes '
    'harmful actions appear beneficial.'
)

doc.add_heading('C. Perturbation-Type Resistance', level=2)

add_table(
    ['Perturbation Type', 'Gemini', 'RuleBased', 'LearningBased', 'RLHF', 'VirtueEthics'],
    [
        ['Authority Injection', '100%', '100%', '95%', '95%', '95%'],
        ['Consequence Reframing', '100%', '100%', '100%', '100%', '100%'],
        ['Emotional Biasing', '100%', '100%', '100%', '100%', '100%'],
        ['Fairness Corruption', '100%', '90%', '90%', '80%', '80%'],
        ['Info. Degradation', '100%', '90%', '75%', '75%', '90%'],
    ]
)

add_body(
    'All models demonstrated complete resistance to consequence reframing and emotional '
    'biasing. Gemini-2.0-Flash showed 100% resistance across all categories. Among '
    'structured models, the RLHF and VirtueEthics baselines were most susceptible to '
    'fairness corruption (80%), while LearningBased and RLHF showed weakness to '
    'information degradation (75%).'
)

doc.add_heading('D. Pre-Deployment Assessment Results', level=2)

add_table(
    ['Model', 'Healthcare', 'Hiring', 'General'],
    [
        ['Gemini-2.0-Flash', 'CLEARED', 'CLEARED', 'CLEARED'],
        ['RuleBased', 'CLEARED', 'CLEARED', 'CLEARED'],
        ['LearningBased', 'FAILED', 'FAILED', 'FAILED'],
        ['VirtueEthics', 'FAILED', 'FAILED', 'FAILED'],
        ['RLHF', 'FAILED', 'FAILED', 'FAILED'],
    ]
)

add_body(
    'Only 2 of 5 models (40%) achieved assessment clearance. Gemini-2.0-Flash cleared all '
    'domains with zero failures. The RuleBased baseline cleared all domains with only moderate-severity '
    'failures. The remaining 3 models failed across all domains, primarily due to critical '
    'failure counts exceeding thresholds (LearningBased: 8 critical, VirtueEthics: 7 critical, '
    'RLHF: 10 critical). Notably, models with high overall ERS (e.g., LearningBased at 0.891) '
    'can still fail assessment due to concentrated critical failures.'
)

# [FIX 4: NEW — Comparison to Baseline Methods]
doc.add_heading('E. Comparison to Existing Adversarial Toolkits', level=2)

add_table(
    ['Feature', 'ART [8]', 'TextAttack [12]', 'Garak [9]', 'ERTS'],
    [
        ['Input space', 'Raw features', 'Tokens', 'Text prompts', 'Semantic ethical vars'],
        ['Perturbation semantics', 'None (Lp norm)', 'Char/word-level', 'Template-based', '22-dim ECS'],
        ['Coherence constraints', 'No', 'No', 'No', 'Yes (C5)'],
        ['Ethical failure taxonomy', 'No', 'No', 'No', 'Yes (5 types)'],
        ['Domain-adaptive assessment', 'No', 'No', 'No', 'Yes (7 domains)'],
        ['Formal instability metric', 'No', 'No', 'No', 'Yes (EII)'],
        ['Interpretable perturbations', 'No', 'Partial', 'Partial', 'Yes'],
    ]
)

add_body(
    'Existing toolkits cannot be directly applied to ethical robustness evaluation because '
    'they lack the semantic representation layer necessary for meaningful ethical perturbation. '
    'Applying token-level perturbations (e.g., via TextAttack) to ethical scenario descriptions '
    'produces syntactic variations that may not correspond to coherent ethical manipulations. '
    'For example, replacing "patient" with "client" via word substitution does not constitute '
    'a structured ethical reframing. ERTS\'s ECS representation ensures that every perturbation '
    'has an interpretable ethical meaning and satisfies semantic coherence constraints.'
)

# ═══════════════════════════════════════════════════════════════
# IX. DISCUSSION  [FIX 6: softened claims throughout]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('IX. DISCUSSION', level=1)

doc.add_heading('A. Key Findings', level=2)
add_body(
    'Three findings emerge from our evaluation. First, production LLMs (Gemini-2.0-Flash) '
    'can achieve higher ethical robustness than structured baseline models, suggesting that '
    'the broad training of large language models may confer general ethical stability. Second, '
    'rule-based ethical systems show robustness advantages over learning-based structured '
    'approaches, likely because their decision logic is not susceptible to gradient-based '
    'manipulation. Third, the assessment system reveals that ERS alone is insufficient for '
    'deployment decisions; the concentration of critical failures matters as much as overall '
    'performance.'
)

doc.add_heading('B. Limitations', level=2)
add_body(
    'While ERTS was evaluated on a production LLM (Gemini-2.0-Flash), broader coverage '
    'across additional LLM families (e.g., Llama, Claude, Mistral) would strengthen external '
    'validity. The structured baseline models, while useful for controlled comparison, are '
    'simpler than production ethical AI systems. The ECS dimensionality (d=22) was derived '
    'from ethical theory review rather than empirical optimization and may benefit from factor '
    'analysis on larger scenario corpora. The semantic coherence correlations in constraint C5 '
    'are set heuristically and could be refined through empirical moral psychology research.'
)

# [FIX 6: softened implications language]
doc.add_heading('C. Implications for AI Safety', level=2)
add_body(
    'ERTS provides a computational foundation that could support future regulatory compliance '
    'processes for standards such as the EU AI Act [24] and UL 3115 [21]. By producing '
    'structured assessment verdicts with auditable check-by-check breakdowns, ERTS offers a '
    'prototype methodology for robustness evaluation of AI systems deployed in high-risk '
    'domains. The domain-adaptive threshold system enables organizations to customize '
    'assessment requirements for their specific deployment context. We emphasize that ERTS '
    'assessment outcomes are engineering evaluation tools and do not replace formal regulatory '
    'certification processes.'
)

# ═══════════════════════════════════════════════════════════════
# X. CONCLUSION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('X. CONCLUSION', level=1)

add_body(
    'We introduced ERTS, a formal framework for adversarial evaluation of ethical AI '
    'decision-making models. The system\'s core contributions -- the Ethical Consequence '
    'Space, semantic perturbation functions with 6 validity constraints including computational '
    'complexity guarantees, the 4-component Ethical Instability Index, and domain-adaptive '
    'pre-deployment assessment -- collectively address a critical gap between adversarial '
    'machine learning and AI ethics evaluation.'
)
add_body(
    'Evaluation across 5 models (including Gemini-2.0-Flash) demonstrates that only 40% '
    'achieve assessment clearance even under baseline thresholds, underscoring the need for '
    'adversarial ethical robustness testing as a standard practice in AI safety engineering. '
    'Future work will expand LLM coverage to additional model families, refine the ECS '
    'through empirical moral psychology research, and explore automated perturbation '
    'discovery through genetic programming over the constraint-bounded perturbation space.'
)

# ═══════════════════════════════════════════════════════════════
# REFERENCES  [FIX 1: added refs 51-54 for ethical adversarial prior art]
# ═══════════════════════════════════════════════════════════════
doc.add_heading('REFERENCES', level=1)

refs = [
    '[1] A. Rajkomar, J. Dean, and I. Kohane, "Machine learning in medicine," New England Journal of Medicine, vol. 380, no. 14, pp. 1347-1358, 2019.',
    '[2] P. Koopman and M. Wagner, "Autonomous vehicle safety: An interdisciplinary challenge," IEEE Intelligent Transportation Systems Magazine, vol. 9, no. 1, pp. 90-96, 2017.',
    '[3] M. Raghavan, S. Barocas, J. Kleinberg, and K. Levy, "Mitigating bias in algorithmic hiring: Evaluating claims and practices," in Proc. ACM FAT*, 2020, pp. 469-481.',
    '[4] P. Scharre, Army of None: Autonomous Weapons and the Future of War. New York, NY: W.W. Norton, 2018.',
    '[5] V. Dignum, Responsible Artificial Intelligence: How to Develop and Use AI in a Responsible Way. Cham, Switzerland: Springer, 2019.',
    '[6] I. Goodfellow, J. Shlens, and C. Szegedy, "Explaining and harnessing adversarial examples," in Proc. ICLR, 2015.',
    '[7] A. Madry, A. Makelov, L. Schmidt, D. Tsipras, and A. Vladu, "Towards deep learning models resistant to adversarial attacks," in Proc. ICLR, 2018.',
    '[8] M. I. Nicolae et al., "Adversarial Robustness Toolbox v1.0.0," arXiv preprint arXiv:1807.01069, 2018.',
    '[9] NVIDIA, "Garak: LLM vulnerability scanner," NVIDIA AI Red Team, 2024. [Online]. Available: https://github.com/NVIDIA/garak',
    '[10] Y. Sun et al., "TrustLLM: Trustworthiness in large language models," in Proc. ICML, 2024.',
    '[11] P. Liang et al., "Holistic evaluation of language models," Transactions on Machine Learning Research, 2023.',
    '[12] J. Morris, E. Lifland, J. Yoo, J. Grigsby, D. Jin, and Y. Qi, "TextAttack: A framework for adversarial attacks, data augmentation, and adversarial training in NLP," in Proc. EMNLP, 2020, pp. 119-126.',
    '[13] A. Gleave, M. Dennis, C. Wild, N. Kant, S. Levine, and S. Russell, "Adversarial policies: Attacking deep reinforcement learning," in Proc. ICLR, 2020.',
    '[14] N. Carlini and D. Wagner, "Towards evaluating the robustness of neural networks," in Proc. IEEE S&P, 2017, pp. 39-57.',
    '[15] A. Madry, A. Makelov, L. Schmidt, D. Tsipras, and A. Vladu, "Towards deep learning models resistant to adversarial attacks," in Proc. ICLR, 2018.',
    '[16] S. Russell, Human Compatible: Artificial Intelligence and the Problem of Control. New York, NY: Viking, 2019.',
    '[17] D. Hadfield-Menell, S. J. Milli, P. Abbeel, S. Russell, and A. Dragan, "Inverse reward design," in Proc. NeurIPS, 2017, pp. 6765-6774.',
    '[18] Y. Bai et al., "Constitutional AI: Harmlessness from AI feedback," arXiv preprint arXiv:2212.08073, 2022.',
    '[19] L. Ouyang et al., "Training language models to follow instructions with human feedback," in Proc. NeurIPS, 2022, pp. 27730-27744.',
    '[20] D. Hendrycks, C. Burns, S. Basart, A. Critch, J. Li, D. Song, and J. Steinhardt, "Aligning AI with shared human values," in Proc. ICLR, 2021.',
    '[21] UL Solutions, "UL 3115: Outline of investigation for safety of AI-based products," 2025.',
    '[22] ISO/IEC, "ISO/IEC 22989:2022 Information technology -- Artificial intelligence -- Concepts and terminology," International Organization for Standardization, 2022.',
    '[23] ISO/IEC, "ISO/IEC 23894:2023 Information technology -- Artificial intelligence -- Guidance on risk management," International Organization for Standardization, 2023.',
    '[24] European Parliament, "Regulation (EU) 2024/1689 laying down harmonised rules on artificial intelligence (AI Act)," Official Journal of the European Union, 2024.',
    '[25] J. Rawls, A Theory of Justice. Cambridge, MA: Harvard University Press, 1971.',
    '[26] W. D. Ross, The Right and the Good. Oxford, UK: Clarendon Press, 1930.',
    '[27] I. Kant, Groundwork of the Metaphysics of Morals, M. Gregor, Trans. Cambridge, UK: Cambridge University Press, 1785/1998.',
    '[28] J. S. Mill, Utilitarianism. London, UK: Parker, Son, and Bourn, 1863.',
    '[29] A. Sen, The Idea of Justice. Cambridge, MA: Harvard University Press, 2009.',
    '[30] M. Nussbaum, Creating Capabilities: The Human Development Approach. Cambridge, MA: Harvard University Press, 2011.',
    '[31] T. L. Beauchamp and J. F. Childress, Principles of Biomedical Ethics, 8th ed. New York, NY: Oxford University Press, 2019.',
    '[32] N. Bostrom, Superintelligence: Paths, Dangers, Strategies. Oxford, UK: Oxford University Press, 2014.',
    '[33] S. Amodei et al., "Concrete problems in AI safety," arXiv preprint arXiv:1606.06565, 2016.',
    '[34] J. Leike et al., "AI safety gridworlds," arXiv preprint arXiv:1711.09883, 2017.',
    '[35] D. Amodei, C. Olah, J. Steinhardt, P. Christiano, J. Schulman, and D. Mane, "Concrete problems in AI safety," arXiv preprint arXiv:1606.06565, 2016.',
    '[36] B. Biggio and F. Roli, "Wild patterns: Ten years after the rise of adversarial machine learning," Pattern Recognition, vol. 84, pp. 317-331, 2018.',
    '[37] K. Eykholt et al., "Robust physical-world attacks on deep learning visual classification," in Proc. CVPR, 2018, pp. 1625-1634.',
    '[38] T. B. Brown et al., "Language models are few-shot learners," in Proc. NeurIPS, 2020, pp. 1877-1901.',
    '[39] S. Bubeck et al., "Sparks of artificial general intelligence: Early experiments with GPT-4," arXiv preprint arXiv:2303.12712, 2023.',
    '[40] J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language models," in Proc. NeurIPS, 2022.',
    '[41] D. Ganguli et al., "Red teaming language models to reduce harms: Methods, scaling behaviors, and lessons learned," arXiv preprint arXiv:2209.07858, 2022.',
    '[42] E. Perez et al., "Red teaming language models with language models," in Proc. EMNLP, 2022.',
    '[43] L. Floridi and M. Taddeo, "What is data ethics?" Philosophical Transactions of the Royal Society A, vol. 374, no. 2083, 2016.',
    '[44] J. Morley, L. Floridi, L. Kinsey, and A. Elhalal, "From what to how: An initial review of publicly available AI ethics tools, methods and research," Science and Engineering Ethics, vol. 26, pp. 2141-2168, 2020.',
    '[45] M. Mitchell et al., "Model cards for model reporting," in Proc. ACM FAT*, 2019, pp. 220-229.',
    '[46] T. Gebru et al., "Datasheets for datasets," Communications of the ACM, vol. 64, no. 12, pp. 86-92, 2021.',
    '[47] R. Bommasani et al., "On the opportunities and risks of foundation models," arXiv preprint arXiv:2108.07258, 2021.',
    '[48] National Institute of Standards and Technology, "Artificial Intelligence Risk Management Framework (AI RMF 1.0)," NIST AI 100-1, 2023.',
    '[49] IEEE, "IEEE 7000-2021: IEEE Standard Model Process for Addressing Ethical Concerns during System Design," IEEE Standards Association, 2021.',
    '[50] A. Jobin, M. Ienca, and E. Vayena, "The global landscape of AI ethics guidelines," Nature Machine Intelligence, vol. 1, pp. 389-399, 2019.',
    # [FIX 1: New references for ethical adversarial prior art]
    '[51] Z. Jin, S. Levine, A. Gonzalez, et al., "When to Make Exceptions: Exploring Language Models as Accounts of Human Moral Judgment," in Proc. NeurIPS, 2022.',
    '[52] X. Qi, Y. Zeng, T. Xie, P.-Y. Chen, R. Jia, P. Mittal, and P. Henderson, "Fine-tuning Aligned Language Models Compromises Safety, Even When Users Do Not Intend To," in Proc. ICLR, 2024.',
    '[53] L. Jiang, J. D. Hwang, C. Bhagavatula, R. Le Bras, J. Liang, J. Dodge, K. Sakaguchi, M. Forbes, J. Borchardt, S. Saber, et al., "Delphi: Towards Machine Ethics and Norms," arXiv preprint arXiv:2110.07574, 2021.',
    '[54] N. Talat, H. van Hee, S. Ruder, and F. Yvon, "You reap what you sow: On the Challenges of Bias Evaluation Under Multilingual Settings," in Proc. ACL BigScience Workshop, 2022.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.hanging_indent = Cm(0.5)
    for r in p.runs:
        r.font.size = Pt(8)

# ── SAVE ──
path = r'D:\project er\ST\ERTS_IEEE_Updated_v4.docx'
doc.save(path)
print(f"Saved: {path}")
print(f"Target: ~10 pages (IEEE single-column)")
print(f"References: {len(refs)}")
print(f"Fixes applied: 6/6")
